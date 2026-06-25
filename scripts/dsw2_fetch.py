#!/usr/bin/env python3
"""dsw2 Fáze 3A — lokální stažení a konverze dokumentů odkazovaných z výzev/programů.

Čte data/dsw2_links.jsonl (z extract/dsw2.py). Dokumenty (kind=doc) stáhne a
převede na text: ODT/DOC/DOCX/RTF/XLS přes `textutil` (macOS), PDF přes `pdftotext`.
Externí webové stránky (kind=web) se zde NEstahují — vypíše jejich seznam do
souboru pro samostatný Apify krok (Fáze 3B), protože arbitrary weby řeší Apify.

Výstup:
  data/dsw2_files/<host>/<sha8>.<ext>      stažený originál
  data/dsw2_files/<host>/<sha8>.txt        extrahovaný text
  data/dsw2_documents.jsonl                manifest {url, soubor, sha256, bytes, chars, status}
  data/dsw2_web_urls.json                  unikátní externí weby pro Apify

Konfigurace přes CLI. Strop velikosti souboru (--max-mb) brání stažení obřích
příloh; whitelist přípon je v extract/dsw2.py (DOC_EXTS).
"""
import argparse
import hashlib
import json
import os
import re
import subprocess
import sys
import urllib.error
import urllib.request
from urllib.parse import urlsplit, urlunsplit, quote
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from limits import L   # centrální registr limitů (root limits.json)
import http_util       # JEDNOTNÁ TLS politika (audit #7/#32: doc-store musí stahovat se stejnou politikou jako harvestery)


def safe_url(url: str) -> str:
    """Percent-encode path+query (mezery, diakritika v názvech souborů) — urllib jinak
    spadne na InvalidURL. safe="%..." nepřekóduje už zakódované URL (žádný double-encode)."""
    p = urlsplit(url)
    path = quote(p.path, safe="/%:@!$&'()*+,;=~-._")
    query = quote(p.query, safe="/%:@!$&'()*+,;=~-._?")
    return urlunsplit((p.scheme, p.netloc, path, query, p.fragment))

UA = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 Chrome/124.0 Safari/537.36"
# Konverzní nástroje dle přípony (macOS): textutil pro Office/ODF, pdftotext pro PDF,
# openpyxl/xlrd pro Excel. POZOR: textutil Excel NEUMÍ (xls/xlsx by přečetl jako syrové byty
# → balast). Proto xls/xlsx jdou samostatnou větví (sešity bývají multi-sheet → všechny listy).
TEXTUTIL_EXTS = {"doc", "docx", "odt", "rtf", "ppt", "pptx"}
SHEET_EXTS = {"xls", "xlsx"}
DOC_EXTS = TEXTUTIL_EXTS | SHEET_EXTS | {"pdf"}
DOC_EXT_RE = re.compile(r"\.(pdf|docx?|rtf|odt|ods|xlsx?|pptx?)(?:$|\?)", re.I)

# MIME → přípona: rozpozná dokument schovaný za URL bez přípony (ASP.NET File.ashx
# handler radnic, ?id_dokumenty=… apod.) z Content-Type / Content-Disposition.
MIME_EXT = {
    "application/pdf": "pdf",
    "application/msword": "doc",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xls",
    "application/vnd.oasis.opendocument.text": "odt",
    "application/vnd.oasis.opendocument.spreadsheet": "ods",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation": "pptx",
    "application/rtf": "rtf", "text/rtf": "rtf",
}


def sniff_ext(url: str, timeout: int):
    """Zjistí příponu dokumentu z hlaviček (Content-Type / Content-Disposition).
    Vrací příponu (str) když jde o dokument, jinak None (= je to webová stránka)."""
    try:
        req = urllib.request.Request(safe_url(url), headers={"User-Agent": UA})
        with http_util.urlopen(req, timeout=timeout) as r:
            ctype = (r.headers.get("Content-Type") or "").split(";")[0].strip().lower()
            cdisp = r.headers.get("Content-Disposition") or ""
        if ctype in MIME_EXT:
            return MIME_EXT[ctype]
        m = re.search(r'filename\*?=(?:UTF-8\'\')?["\']?[^"\';]*\.(\w{2,5})', cdisp, re.I)
        if m and m.group(1).lower() in DOC_EXTS:
            return m.group(1).lower()
    except Exception:  # noqa: BLE001
        pass
    return None


def host_of(u: str) -> str:
    m = re.match(r"https?://([^/]+)", u)
    return re.sub(r"[^a-z0-9.]+", "_", m.group(1).lower()) if m else "unknown"


def ext_of(u: str) -> str:
    m = re.search(r"\.(\w{2,5})(?:$|\?)", u)
    return m.group(1).lower() if m else "bin"


def download(url: str, dest: str, timeout: int, max_bytes: int):
    try:
        req = urllib.request.Request(safe_url(url), headers={"User-Agent": UA})
        with http_util.urlopen(req, timeout=timeout) as r:
            data = r.read(max_bytes + 1)
        if len(data) > max_bytes:
            return None, "too-big"
        with open(dest, "wb") as f:
            f.write(data)
        return len(data), None
    except urllib.error.HTTPError as e:
        return None, f"http-{e.code}"
    except Exception as e:  # noqa: BLE001
        return None, f"{type(e).__name__}: {str(e)[:60]}"


def _spreadsheet_text(path: str, ext: str) -> str:
    """xls/xlsx → text VŠECH listů (textutil Excel neumí). openpyxl pro xlsx, xlrd pro xls."""
    out = []
    if ext == "xlsx":
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        for ws in wb.worksheets:
            out.append(f"## List: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if any(c.strip() for c in cells):
                    out.append("\t".join(cells).rstrip())
        wb.close()
    else:  # xls
        import xlrd
        wb = xlrd.open_workbook(path)
        for sh in wb.sheets():
            out.append(f"## List: {sh.name}")
            for r in range(sh.nrows):
                cells = [("" if v == "" else str(v)) for v in sh.row_values(r)]
                if any(str(c).strip() for c in cells):
                    out.append("\t".join(cells).rstrip())
    return "\n".join(out)


def _docx_text(path: str) -> str:
    """docx → text PLATFORMOVĚ NEZÁVISLE (python-docx). Pokryje i soubory s příponou .doc,
    které jsou ve skutečnosti OOXML (časté u MŠMT). Audit #10: textutil je jen macOS."""
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def convert(path: str, ext: str, txt_path: str, timeout: int):
    """Vrátí (chars:int|None, err:str|None) a zapíše text do txt_path."""
    try:
        if ext == "pdf":
            subprocess.run(["pdftotext", "-q", path, txt_path],
                           timeout=timeout, check=False)
            text = open(txt_path, encoding="utf-8", errors="replace").read() \
                if os.path.exists(txt_path) else ""
        elif ext in SHEET_EXTS:                       # Excel: openpyxl/xlrd, NE textutil
            text = _spreadsheet_text(path, ext)
            open(txt_path, "w", encoding="utf-8").write(text)
        elif ext == "docx":                           # cross-platform (python-docx); fallback textutil (macOS)
            try:
                text = _docx_text(path)
            except Exception:
                r = subprocess.run(["textutil", "-convert", "txt", "-stdout", path],
                                   capture_output=True, timeout=timeout)
                text = r.stdout.decode("utf-8", "replace")
            open(txt_path, "w", encoding="utf-8").write(text)
        elif ext in TEXTUTIL_EXTS:                    # doc/odt/rtf/ppt/pptx → textutil (macOS-only, audit #10)
            r = subprocess.run(["textutil", "-convert", "txt", "-stdout", path],
                               capture_output=True, timeout=timeout)
            text = r.stdout.decode("utf-8", "replace")
            open(txt_path, "w", encoding="utf-8").write(text)
        else:
            return None, f"no-converter-for-{ext}"
        return len(text), None
    except Exception as e:  # noqa: BLE001
        return None, f"convert: {type(e).__name__}: {str(e)[:50]}"


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--links", default="data/dsw2_links.jsonl")
    ap.add_argument("--dataset", default="data/merged_dataset.json")
    ap.add_argument("--files-dir", default="data/dsw2_files")
    ap.add_argument("--out-docs", default="data/dsw2_documents.jsonl")
    ap.add_argument("--out-web", default="data/dsw2_web_urls.json")
    ap.add_argument("--timeout", type=int, default=40)
    ap.add_argument("--max-mb", type=int, default=L("safety.doc_download_max_mb"),
                    help="runaway-pojistka velikosti souboru (limits.json safety.doc_download_max_mb)")
    args = ap.parse_args()

    links = [json.loads(l) for l in open(args.links, encoding="utf-8")]
    # dsw2 hostitelé = interní (project_detail apod. neřešíme tady)
    bases = {re.match(r"https?://([^/]+)", r["url"]).group(1)
             for r in json.load(open(args.dataset, encoding="utf-8"))
             if r.get("platform") == "dsw2_otevrena_mesta"}

    # mapa url -> entity (kvůli manifestu)
    ents = {}
    for l in links:
        ents.setdefault(l["url"], []).append((l["entity_type"], l.get("entity_title")))

    # dokumenty dle přípony v URL
    to_download = {l["url"]: ext_of(l["url"]) for l in links if l["kind"] == "doc"}
    # externí webové odkazy → očichat: dokument za handlerem (File.ashx) → stáhnout
    web_cands = sorted({l["url"].replace("&amp;", "&").strip()
                        for l in links if l["kind"] == "web"
                        and l["entity_type"] in ("appeal", "program")
                        and re.match(r"https?://([^/]+)", l["url"]).group(1) not in bases})
    web = []
    for url in web_cands:
        ext = sniff_ext(url, args.timeout) if not DOC_EXT_RE.search(url) else ext_of(url)
        if ext:
            to_download[url] = ext
            print(f"  [sniff] dokument za URL bez přípony → .{ext}: {url[:60]}", file=sys.stderr)
        else:
            web.append(url)

    os.makedirs(args.files_dir, exist_ok=True)
    max_bytes = args.max_mb * 1024 * 1024
    manifest = []
    for url, ext in sorted(to_download.items()):
        h = host_of(url)
        sha = hashlib.sha256(url.encode()).hexdigest()[:8]
        ddir = os.path.join(args.files_dir, h)
        os.makedirs(ddir, exist_ok=True)
        fpath = os.path.join(ddir, f"{sha}.{ext}")
        tpath = os.path.join(ddir, f"{sha}.txt")
        nbytes, derr = download(url, fpath, args.timeout, max_bytes)
        chars = cerr = None
        if nbytes:
            sha256 = hashlib.sha256(open(fpath, "rb").read()).hexdigest()
            chars, cerr = convert(fpath, ext, tpath, args.timeout)
        else:
            sha256 = None
        rec = {"url": url, "entities": ents.get(url, []), "ext": ext,
               "file": fpath if nbytes else None, "text_file": tpath if chars else None,
               "bytes": nbytes, "sha256": sha256, "text_chars": chars,
               "status": "ok" if chars else (derr or cerr or "fail")}
        manifest.append(rec)
        print(f"  {rec['status']:14} {ext:4} chars={str(chars):>6} {url[:70]}", file=sys.stderr)

    with open(args.out_docs, "w", encoding="utf-8") as f:
        for r in manifest:
            f.write(json.dumps(r, ensure_ascii=False) + "\n")
    json.dump(web, open(args.out_web, "w", encoding="utf-8"), ensure_ascii=False, indent=1)

    ok = sum(1 for r in manifest if r["status"] == "ok")
    print(json.dumps({"MARKER": "DSW2_FETCH", "docs_total": len(to_download), "docs_ok": ok,
                      "web_urls_for_apify": len(web), "out_docs": args.out_docs,
                      "out_web": args.out_web}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
